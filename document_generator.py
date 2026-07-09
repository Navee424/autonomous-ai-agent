from docx import Document
from docx.shared import Pt
from datetime import datetime
import os


def generate_word_document(content: str, title: str = "Business Proposal") -> str:
    """
    Generates a Microsoft Word document.
    Returns the generated file path.
    """

    try:

        os.makedirs("generated_docs", exist_ok=True)

        document = Document()

        document.add_heading(title, level=1)

        document.add_paragraph(
            f"Generated On: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}"
        )

        document.add_paragraph()

        paragraph = document.add_paragraph()

        run = paragraph.add_run(content)

        run.font.size = Pt(11)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        safe_title = title.replace(" ", "_")

        filename = f"generated_docs/{safe_title}_{timestamp}.docx"

        document.save(filename)

        return filename

    except Exception as e:

        raise RuntimeError(
            f"Failed to generate Word document: {e}"
        )